import subprocess
import tempfile
from pathlib import Path

from django.core.files import File
from django.utils.text import slugify
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from apps.specs.instructions import instruction_for_unit
from apps.specs.models import DocumentExport, TechnicalSpec

REQUIREMENTS_LINE = (
    "Требования заказчика к объекту закупки по техническим, функциональным и качественным характеристикам, "
    "эксплуатационным характеристикам объекта закупки"
)

DOCX_FONT_NAME = "Times New Roman"
DOCX_FONT_SIZE = Pt(10)


def _safe_name(spec: TechnicalSpec, fmt: str) -> str:
    base = slugify(spec.title, allow_unicode=True) or f"spec-{spec.id}"
    return f"{base}-{spec.id}.{fmt}"


def _included_characteristics(item):
    return item.selected_characteristics.filter(is_active=True).order_by("display_order", "id")


def _postscript_lines(spec: TechnicalSpec) -> list[str]:
    template_lines = [template.text.strip() for template in spec.postscript_templates.filter(is_active=True).order_by("name", "id") if template.text.strip()]
    custom_lines = [line.strip() for line in spec.custom_postscript.splitlines() if line.strip()]
    return template_lines + custom_lines


def _format_quantity(value) -> str:
    return str(value).rstrip("0").rstrip(".")


def _apply_run_font(run, size=DOCX_FONT_SIZE) -> None:
    run.font.name = DOCX_FONT_NAME
    run.font.size = size
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), DOCX_FONT_NAME)


def _apply_style_font(style, size=DOCX_FONT_SIZE) -> None:
    style.font.name = DOCX_FONT_NAME
    style.font.size = size
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), DOCX_FONT_NAME)


def _compact_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)


def _set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    _compact_paragraph(paragraph)
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    _apply_run_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _shade_cell(cell, fill: str) -> None:
    return None


def _set_cell_margins(cell, top=0, start=25, bottom=0, end=25) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_row_auto_height(row) -> None:
    row.height = None
    row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
    tr_pr = row._tr.get_or_add_trPr()
    for tr_height in tr_pr.findall(qn("w:trHeight")):
        tr_pr.remove(tr_height)


def _format_table(table, widths: list[float], header_rows: int = 1) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for index, width in enumerate(widths):
        for cell in table.columns[index].cells:
            cell.width = Cm(width)
            _set_cell_margins(cell)
    for row_index, row in enumerate(table.rows):
        _set_row_auto_height(row)
        if row_index < header_rows:
            _repeat_table_header(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _add_postscripts(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    table = document.add_table(rows=len(lines), cols=1)
    _format_table(table, [15.8], header_rows=0)
    for index, line in enumerate(lines, start=1):
        _set_cell_text(table.cell(index - 1, 0), f"{index}. {line}")


def build_docx(spec: TechnicalSpec, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    if spec.items.exists():
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    styles = document.styles
    _apply_style_font(styles["Normal"])
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    styles["Normal"].paragraph_format.line_spacing = 1

    postscript_lines = _postscript_lines(spec)

    for item_index, item in enumerate(spec.items.select_related("ktru_position").prefetch_related("selected_characteristics"), start=1):
        if item_index > 1:
            document.add_page_break()

        top = document.add_table(rows=2, cols=5)
        _format_table(top, [1.1, 8.2, 2.8, 2.2, 1.5])
        headers = ["N п/п", "Наименование объекта закупки", "Единица измерения", "Кол-во", "КТРУ"]
        values = [
            str(item.position_number),
            item.object_name,
            item.unit_name,
            _format_quantity(item.quantity),
            item.ktru_position.code,
        ]
        for value, cell in zip(headers, top.rows[0].cells):
            _set_cell_text(cell, value, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade_cell(cell, "1F2937")
        for value, cell in zip(values, top.rows[1].cells):
            _set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER if cell in (top.rows[1].cells[0], top.rows[1].cells[2], top.rows[1].cells[3]) else WD_ALIGN_PARAGRAPH.LEFT)

        meta = document.add_table(rows=1, cols=2)
        _format_table(meta, [3.1, 12.7])
        _set_cell_text(meta.cell(0, 0), "ОКПД-2", bold=True)
        _shade_cell(meta.cell(0, 0), "E5E7EB")
        _set_cell_text(meta.cell(0, 1), item.ktru_position.okpd2_code or "Не указан")

        line_table = document.add_table(rows=1, cols=1)
        _format_table(line_table, [15.8], header_rows=0)
        _set_cell_text(line_table.cell(0, 0), REQUIREMENTS_LINE, bold=True)
        _shade_cell(line_table.cell(0, 0), "EFF6FF")

        characteristics = list(_included_characteristics(item))
        rows_count = max(len(characteristics) + 1, 2)
        table = document.add_table(rows=rows_count, cols=5)
        _format_table(table, [2.8, 4.4, 3.6, 1.8, 3.2])
        headers = [
            "Объект закупки",
            "Наименование характеристики",
            "Значение характеристики",
            "Единица измерения",
            "Инструкция по заполнению",
        ]
        for value, cell in zip(headers, table.rows[0].cells):
            _set_cell_text(cell, value, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade_cell(cell, "1F2937")

        if characteristics:
            table.cell(1, 0).text = item.object_name
            table.cell(1, 0).merge(table.cell(len(characteristics), 0))
            _set_cell_text(table.cell(1, 0), item.object_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _set_cell_text(table.cell(1, 0), item.object_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for row_index, characteristic in enumerate(characteristics, start=1):
            row = table.rows[row_index]
            _set_cell_text(row.cells[1], characteristic.characteristic_name_snapshot)
            _set_cell_text(row.cells[2], characteristic.display_value)
            _set_cell_text(row.cells[3], characteristic.unit_name_snapshot, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[4], instruction_for_unit(characteristic.unit_name_snapshot, characteristic.instruction_snapshot))
        _add_postscripts(document, postscript_lines)

    document.save(path)


def build_xlsx(spec: TechnicalSpec, path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "ТЗ"
    headers = [
        "N позиции",
        "Название ТЗ",
        "Наименование объекта закупки",
        "КТРУ",
        "ОКПД-2",
        "Единица измерения позиции",
        "Количество",
        "Наименование характеристики",
        "Значение характеристики",
        "Единица измерения характеристики",
        "Инструкция по заполнению характеристики в заявке",
        "Обязательная/необязательная",
        "Активная/исключенная",
    ]
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F2937")
        cell.alignment = Alignment(wrap_text=True, vertical="center")

    for item in spec.items.select_related("ktru_position").prefetch_related("selected_characteristics"):
        for characteristic in _included_characteristics(item):
            sheet.append(
                [
                    item.position_number,
                    spec.title,
                    item.object_name,
                    item.ktru_position.code,
                    item.ktru_position.okpd2_code,
                    item.unit_name,
                    float(item.quantity),
                    characteristic.characteristic_name_snapshot,
                    ", ".join(characteristic.selected_values),
                    characteristic.unit_name_snapshot,
                    instruction_for_unit(characteristic.unit_name_snapshot, characteristic.instruction_snapshot),
                    "обязательная" if characteristic.is_required_snapshot else "необязательная",
                    "активная" if characteristic.is_active else "исключенная",
                ]
            )

    widths = [12, 28, 32, 24, 18, 18, 12, 42, 30, 22, 56, 24, 20]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    sheet.freeze_panes = "A2"
    workbook.save(path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    output_dir = pdf_path.parent
    result = subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=90,
    )
    generated = output_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not generated.exists():
        raise RuntimeError(f"LibreOffice не смог сформировать PDF: {result.stderr or result.stdout}")
    generated.replace(pdf_path)


def export_spec(spec: TechnicalSpec, fmt: str, user):
    fmt = fmt.lower()
    if fmt not in {"docx", "xlsx", "pdf"}:
        raise ValueError("Unsupported export format")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        if fmt == "docx":
            output = tmp_dir / _safe_name(spec, "docx")
            build_docx(spec, output)
        elif fmt == "xlsx":
            output = tmp_dir / _safe_name(spec, "xlsx")
            build_xlsx(spec, output)
        else:
            docx = tmp_dir / _safe_name(spec, "docx")
            output = tmp_dir / _safe_name(spec, "pdf")
            build_docx(spec, docx)
            convert_docx_to_pdf(docx, output)

        export = DocumentExport.objects.create(technical_spec=spec, format=fmt, created_by=user)
        with output.open("rb") as fh:
            export.file.save(_safe_name(spec, fmt), File(fh), save=True)
        return export
